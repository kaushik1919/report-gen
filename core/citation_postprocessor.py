"""
Post-process an assembled DOCX to resolve [REF: key] citation markers.

Flow:
1. Scan all paragraphs for [REF: key] markers (produced by the assembler).
2. Assign sequential numbers to unique keys in order of first appearance.
3. Replace each marker with its inline reference: [N] or [?key] if unknown.
4. Render a formatted bibliography and insert it at the {{REFERENCES}} anchor.
   If the anchor is absent, the bibliography is appended at the end of the doc.

No LLM calls are made. All rendering is deterministic.
"""
import logging
import re

from docx import Document

from core.models import BibliographyStore, CitationEntry

logger = logging.getLogger(__name__)

_MARKER_RE = re.compile(r"\[REF:\s*([^\]]+?)\s*\]")
_REFERENCES_ANCHOR = "{{REFERENCES}}"

# BibTeX → CSL type mapping used by _entry_to_csl.
_BIBTEX_TYPE_MAP: dict[str, str] = {
    "article": "article-journal",
    "book": "book",
    "inproceedings": "paper-conference",
    "proceedings": "book",
    "incollection": "chapter",
    "phdthesis": "thesis",
    "mastersthesis": "thesis",
    "techreport": "report",
    "misc": "article",
    "unpublished": "manuscript",
}


def process_document(
    doc: Document,
    bib: BibliographyStore,
) -> list[str]:
    """
    Resolve citation markers in *doc* in-place using *bib*.

    Returns a list of warning strings for any unknown citation keys found.
    """
    # Pass 1 — collect unique keys in order of first appearance.
    seen_keys: list[str] = []
    key_to_number: dict[str, int] = {}
    for para in doc.paragraphs:
        for m in _MARKER_RE.finditer(para.text):
            key = m.group(1).strip()
            if key not in key_to_number:
                key_to_number[key] = len(seen_keys) + 1
                seen_keys.append(key)

    warnings: list[str] = []
    for key in seen_keys:
        if not bib.has(key):
            msg = f"Unknown citation key: {key!r} — no entry in bibliography"
            logger.warning("citation_postprocessor: %s", msg)
            warnings.append(msg)

    if not seen_keys:
        return warnings  # nothing to do

    # Pass 2 — replace marker text in each paragraph.
    for para in doc.paragraphs:
        if _MARKER_RE.search(para.text):
            new_text = _MARKER_RE.sub(
                lambda m: _inline_ref(m.group(1).strip(), key_to_number, bib),
                para.text,
            )
            _set_paragraph_text(para, new_text)

    # Pass 3 — render bibliography and insert at anchor (or append).
    bib_lines = _render_bibliography(seen_keys, key_to_number, bib)

    anchor = _find_anchor(doc, _REFERENCES_ANCHOR)
    if anchor is not None:
        _replace_anchor_with_bibliography(doc, anchor, bib_lines)
    else:
        logger.warning(
            "citation_postprocessor: %r anchor not found — appending bibliography at end",
            _REFERENCES_ANCHOR,
        )
        _append_bibliography(doc, bib_lines)

    return warnings


# ---------------------------------------------------------------------------
# Inline reference formatting
# ---------------------------------------------------------------------------

def _inline_ref(
    key: str,
    key_to_number: dict[str, int],
    bib: BibliographyStore,
) -> str:
    num = key_to_number.get(key)
    if num is None or not bib.has(key):
        return f"[?{key}]"
    return f"[{num}]"


# ---------------------------------------------------------------------------
# Bibliography rendering
# ---------------------------------------------------------------------------

def _render_bibliography(
    keys: list[str],
    key_to_number: dict[str, int],
    bib: BibliographyStore,
) -> list[str]:
    lines: list[str] = []
    for key in keys:
        num = key_to_number[key]
        entry = bib.get(key)
        if entry is None:
            lines.append(f"[{num}] [{key}: reference not found in bibliography]")
        else:
            lines.append(_format_entry(num, entry))
    return lines


def _format_entry(num: int, entry: CitationEntry) -> str:
    """Format one bibliography entry, using citeproc-py if available."""
    try:
        return _format_with_citeproc(num, entry)
    except Exception as exc:
        logger.debug(
            "citation_postprocessor: citeproc-py failed for %r (%s) — using fallback",
            entry.key,
            exc,
        )
        return _format_simple(num, entry)


def _format_with_citeproc(num: int, entry: CitationEntry) -> str:
    """Format using citeproc-py with the bundled chicago-author-date style."""
    import os

    import citeproc as _cp
    from citeproc import (
        Citation,
        CitationItem,
        CitationStylesBibliography,
        CitationStylesStyle,
        formatter,
    )
    from citeproc.source.json import CiteProcJSON

    csl_data = [_entry_to_csl(entry)]
    source = CiteProcJSON(csl_data)

    style_file = os.path.join(
        os.path.dirname(_cp.__file__), "data", "styles", "chicago-author-date.csl"
    )
    style = CitationStylesStyle(style_file, validate=False)
    bibliography = CitationStylesBibliography(style, source, formatter.plain)
    bibliography.register(Citation([CitationItem(entry.key)]))
    bibliography.sort()

    for item in bibliography.bibliography():
        return f"[{num}] {item}"
    return _format_simple(num, entry)


def _format_simple(num: int, entry: CitationEntry) -> str:
    """APA-like fallback formatter that requires no external style files."""
    f = entry.fields
    author = _short_author(f.get("author", ""))
    year = f.get("year", "n.d.")
    title = _strip_braces(f.get("title", entry.key))
    journal = _strip_braces(f.get("journal", "") or f.get("booktitle", ""))
    publisher = _strip_braces(f.get("publisher", ""))

    if journal:
        vol = f.get("volume", "")
        pages = _strip_braces(f.get("pages", ""))
        suffix = f" {journal}"
        if vol:
            suffix += f", {vol}"
        if pages:
            suffix += f", {pages}"
        return f"[{num}] {author} ({year}). {title}.{suffix}."
    if publisher:
        return f"[{num}] {author} ({year}). {title}. {publisher}."
    return f"[{num}] {author} ({year}). {title}."


def _short_author(author_str: str) -> str:
    if not author_str:
        return "Unknown"
    first = author_str.split(" and ")[0].strip()
    if "," in first:
        return first.split(",")[0].strip("{} ")
    if " " in first:
        return first.rsplit(" ", 1)[-1].strip("{} ")
    return first.strip("{} ")


def _strip_braces(text: str) -> str:
    return text.replace("{", "").replace("}", "").strip()


def _entry_to_csl(entry: CitationEntry) -> dict:
    """Convert a CitationEntry to a minimal CSL-JSON dict."""
    csl: dict = {
        "id": entry.key,
        "type": _BIBTEX_TYPE_MAP.get(entry.entry_type, "article"),
    }
    f = entry.fields
    if f.get("title"):
        csl["title"] = _strip_braces(f["title"])
    if f.get("author"):
        csl["author"] = _parse_csl_authors(f["author"])
    year = f.get("year", "")
    if year.isdigit():
        csl["issued"] = {"date-parts": [[int(year)]]}
    container = f.get("journal", "") or f.get("booktitle", "")
    if container:
        csl["container-title"] = _strip_braces(container)
    for bib_field, csl_key in [
        ("volume", "volume"),
        ("number", "issue"),
        ("pages", "page"),
        ("publisher", "publisher"),
    ]:
        if f.get(bib_field):
            csl[csl_key] = _strip_braces(f[bib_field])
    return csl


def _parse_csl_authors(author_str: str) -> list[dict]:
    authors = []
    for a in author_str.split(" and "):
        a = a.strip()
        if "," in a:
            parts = a.split(",", 1)
            authors.append({"family": parts[0].strip("{} "), "given": parts[1].strip("{} ")})
        elif " " in a:
            parts = a.rsplit(" ", 1)
            authors.append({"family": parts[1].strip("{} "), "given": parts[0].strip("{} ")})
        else:
            authors.append({"literal": a.strip("{} ")})
    return authors


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _find_anchor(doc: Document, anchor: str):
    for para in doc.paragraphs:
        if para.text.strip() == anchor:
            return para
    return None


def _replace_anchor_with_bibliography(
    doc: Document, anchor_para, lines: list[str]
) -> None:
    """Replace anchor paragraph with a References heading + bibliography entries."""
    anchor_para.insert_paragraph_before("References", "Heading 1")
    for line in lines:
        anchor_para.insert_paragraph_before(line, "Normal")
    parent = anchor_para._element.getparent()
    if parent is not None:
        parent.remove(anchor_para._element)


def _append_bibliography(doc: Document, lines: list[str]) -> None:
    doc.add_paragraph("References", style="Heading 1")
    for line in lines:
        doc.add_paragraph(line, style="Normal")


def _set_paragraph_text(para, text: str) -> None:
    """Replace paragraph content with a single run containing *text*."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
