import logging
import re
import shutil
from pathlib import Path

from docx import Document

from app.config import OUTPUTS_DIR
from core.models import ReportPlan, SectionContent, TemplateProfile

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

_BODY_ANCHOR = "{{REPORT_BODY}}"
# These anchors are reserved for future milestones; clean them up when unused.
_FUTURE_ANCHORS = frozenset({"{{REFERENCES}}", "{{APPENDIX}}"})

# Canonical block types recognised by the assembler.
# "bullets" is a legacy alias for "bullet_list" and will continue to work.
KNOWN_BLOCK_TYPES = frozenset({
    "paragraph",
    "bullet_list", "bullets",
    "numbered_list",
    "table",
    "heading",
    "figure_placeholder",
    "citation_placeholder",
})


class DocumentAssemblerError(Exception):
    """Raised when document assembly fails."""


def build(
    template_path: Path,
    profile: TemplateProfile,
    plan: ReportPlan,
    sections: list[SectionContent],
) -> Path:
    """
    Merge TemplateProfile + ReportPlan + [SectionContent] into a styled DOCX.
    Returns the path to the assembled output file.

    Anchor behaviour
    ----------------
    If the template contains a paragraph whose text is exactly ``{{REPORT_BODY}}``,
    all section content is inserted at that position and the anchor is removed.
    If the anchor is absent, sections are appended to the end of the document
    with a logged warning.  ``{{REFERENCES}}`` and ``{{APPENDIX}}`` anchors are
    always cleaned up (they are reserved for M4).
    """
    out_path = OUTPUTS_DIR / f"{_slug(plan.title)}.docx"
    shutil.copy2(template_path, out_path)
    doc = Document(str(out_path))

    _replace_placeholders(doc, {"title": plan.title, "author": plan.author})

    body_anchor = _find_anchor(doc, _BODY_ANCHOR)
    if body_anchor is not None:
        _insert_sections_at_anchor(doc, body_anchor, sections, profile)
        parent = body_anchor._element.getparent()
        if parent is not None:
            parent.remove(body_anchor._element)
    else:
        logger.warning("%r anchor not found — appending sections to document end", _BODY_ANCHOR)
        for section in sections:
            _append_section(doc, section, profile)

    _cleanup_future_anchors(doc)
    doc.save(str(out_path))
    logger.info("Assembled document saved to %s", out_path)
    return out_path


# ---------------------------------------------------------------------------
# Anchor helpers
# ---------------------------------------------------------------------------

def _find_anchor(doc: Document, anchor: str):
    for para in doc.paragraphs:
        if para.text.strip() == anchor:
            return para
    return None


def _insert_sections_at_anchor(
    doc: Document, anchor_para, sections: list[SectionContent], profile: TemplateProfile
) -> None:
    """Insert all sections' content as siblings immediately before *anchor_para*."""
    for section in sections:
        heading_style = _resolve_heading_style(section.level, profile)
        anchor_para.insert_paragraph_before(section.title, heading_style)
        for block in section.blocks:
            _insert_block_before(doc, anchor_para, block, profile)


def _insert_block_before(
    doc: Document, anchor_para, block: dict, profile: TemplateProfile
) -> None:
    block_type = block.get("type")
    if block_type not in KNOWN_BLOCK_TYPES:
        logger.warning("Rejecting unknown block type %r — skipping block", block_type)
        return

    body_style = _resolve_body_style(profile)

    if block_type == "paragraph":
        anchor_para.insert_paragraph_before(block.get("text", ""), body_style)
    elif block_type in ("bullet_list", "bullets"):
        style = _resolve_bullet_style(profile)
        for item in block.get("items", []):
            anchor_para.insert_paragraph_before(item, style)
    elif block_type == "numbered_list":
        style = "List Number" if "List Number" in profile.styles else "List Paragraph"
        for item in block.get("items", []):
            anchor_para.insert_paragraph_before(item, style)
    elif block_type == "heading":
        level = int(block.get("level", 2))
        style = _resolve_heading_style(level, profile)
        anchor_para.insert_paragraph_before(block.get("text", ""), style)
    elif block_type == "figure_placeholder":
        caption = block.get("caption", "")
        anchor_para.insert_paragraph_before(f"[Figure: {caption}]", body_style)
    elif block_type == "citation_placeholder":
        key = block.get("key", "")
        anchor_para.insert_paragraph_before(f"[REF: {key}]", body_style)
    elif block_type == "table":
        _insert_table_before(doc, anchor_para, block.get("headers", []), block.get("rows", []))


def _insert_table_before(
    doc: Document, anchor_para, headers: list[str], rows: list[list[str]]
) -> None:
    if not headers:
        return
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row[:col_count]):
            table.rows[r_idx + 1].cells[c_idx].text = cell_text
    # Move table XML element to immediately before the anchor paragraph.
    anchor_para._element.addprevious(table._element)


def _cleanup_future_anchors(doc: Document) -> None:
    """Remove {{REFERENCES}} and {{APPENDIX}} anchor paragraphs if still present."""
    for para in list(doc.paragraphs):
        if para.text.strip() in _FUTURE_ANCHORS:
            logger.info("Removing unused future anchor: %r", para.text.strip())
            parent = para._element.getparent()
            if parent is not None:
                parent.remove(para._element)


# ---------------------------------------------------------------------------
# Fallback section insertion (used when {{REPORT_BODY}} is absent)
# ---------------------------------------------------------------------------

def _append_section(
    doc: Document, section: SectionContent, profile: TemplateProfile
) -> None:
    heading_style = _resolve_heading_style(section.level, profile)
    doc.add_paragraph(section.title, style=heading_style)
    for block in section.blocks:
        _append_block(doc, block, profile)


def _append_block(doc: Document, block: dict, profile: TemplateProfile) -> None:
    block_type = block.get("type")
    if block_type not in KNOWN_BLOCK_TYPES:
        logger.warning("Rejecting unknown block type %r — skipping block", block_type)
        return

    body_style = _resolve_body_style(profile)

    if block_type == "paragraph":
        doc.add_paragraph(block.get("text", ""), style=body_style)
    elif block_type in ("bullet_list", "bullets"):
        style = _resolve_bullet_style(profile)
        for item in block.get("items", []):
            doc.add_paragraph(item, style=style)
    elif block_type == "numbered_list":
        style = "List Number" if "List Number" in profile.styles else "List Paragraph"
        for item in block.get("items", []):
            doc.add_paragraph(item, style=style)
    elif block_type == "heading":
        level = int(block.get("level", 2))
        doc.add_paragraph(block.get("text", ""), style=_resolve_heading_style(level, profile))
    elif block_type == "figure_placeholder":
        doc.add_paragraph(f"[Figure: {block.get('caption', '')}]", style=body_style)
    elif block_type == "citation_placeholder":
        doc.add_paragraph(f"[REF: {block.get('key', '')}]", style=body_style)
    elif block_type == "table":
        _append_table(doc, block.get("headers", []), block.get("rows", []))


def _append_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row[:col_count]):
            table.rows[r_idx + 1].cells[c_idx].text = cell_text


# ---------------------------------------------------------------------------
# Placeholder replacement
# ---------------------------------------------------------------------------

def _replace_placeholders(doc: Document, subs: dict[str, str]) -> None:
    """Replace {{key}} tokens in template paragraphs. Handles cross-run splits."""
    for para in doc.paragraphs:
        if not _PLACEHOLDER_RE.search(para.text):
            continue
        replaced = _PLACEHOLDER_RE.sub(
            lambda m: subs.get(m.group(1), m.group(0)), para.text
        )
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = replaced
        else:
            para.add_run(replaced)


# ---------------------------------------------------------------------------
# Style resolution
# ---------------------------------------------------------------------------

def _resolve_heading_style(level: int, profile: TemplateProfile) -> str:
    name = f"Heading {level}"
    if name in profile.styles:
        return name
    for lvl in range(level - 1, 0, -1):
        candidate = f"Heading {lvl}"
        if candidate in profile.styles:
            logger.warning("Heading %d absent from profile, falling back to %s", level, candidate)
            return candidate
    return name  # python-docx has Heading 1–9 built-in


def _resolve_body_style(profile: TemplateProfile) -> str:
    for candidate in ("Body Text", "Normal"):
        if candidate in profile.styles:
            return candidate
    return "Normal"


def _resolve_bullet_style(profile: TemplateProfile) -> str:
    for candidate in ("List Bullet", "List Paragraph"):
        if candidate in profile.styles:
            return candidate
    return "List Bullet"


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def _slug(text: str) -> str:
    return re.sub(r"[^\w]+", "_", text.lower()).strip("_")[:60]
