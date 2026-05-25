import logging
import re

from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models import LoadedTemplate, StyleSpec, TemplateProfile

logger = logging.getLogger(__name__)

# Matches {{placeholder}} tokens in paragraph text
_PLACEHOLDER_RE = re.compile(r"\{\{([\w\s]+)\}\}")

# Matches "Heading 1", "Heading 2", etc.
_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)

_ALIGNMENT_MAP: dict = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}


def extract(loaded: LoadedTemplate) -> TemplateProfile:
    """
    Deterministically extract all formatting metadata from a loaded DOCX.
    No LLM is involved. Returns a TemplateProfile ready for the AI pipeline.
    """
    doc = loaded.document

    styles = _extract_styles(doc)
    margins, page_size = _extract_page_geometry(doc)
    heading_hierarchy = _extract_heading_hierarchy(doc)
    placeholders = _detect_placeholders(doc)
    section_skeleton = _build_section_skeleton(doc)

    logger.info(
        "Extracted profile: %d styles, %d heading levels, %d placeholders, %d skeleton entries",
        len(styles),
        len(heading_hierarchy),
        len(placeholders),
        len(section_skeleton),
    )

    return TemplateProfile(
        styles=styles,
        margins_in=margins,
        page_size=page_size,
        heading_hierarchy=heading_hierarchy,
        placeholders=placeholders,
        section_skeleton=section_skeleton,
        numbering_scheme=None,
    )


# ---------------------------------------------------------------------------
# Style extraction
# ---------------------------------------------------------------------------

def _extract_styles(doc) -> dict[str, StyleSpec]:
    styles: dict[str, StyleSpec] = {}
    for style in doc.styles:
        if style.type is None:
            continue
        try:
            styles[style.name] = _style_to_spec(style)
        except Exception as exc:
            logger.debug("Skipping style %r: %s", style.name, exc)
    return styles


def _style_to_spec(style) -> StyleSpec:
    font = style.font
    pf = style.paragraph_format

    return StyleSpec(
        name=style.name,
        font_name=font.name or "Default",
        font_size_pt=font.size.pt if font.size else 12.0,
        bold=bool(font.bold),
        italic=bool(font.italic),
        color_hex=_safe_color(font),
        alignment=_safe_alignment(pf),
        line_spacing=_safe_line_spacing(pf),
    )


def _safe_color(font) -> str | None:
    try:
        if font.color and font.color.rgb:
            return str(font.color.rgb)
    except Exception:
        pass
    return None


def _safe_alignment(pf) -> str:
    try:
        if pf.alignment is not None:
            return _ALIGNMENT_MAP.get(pf.alignment, "left")
    except Exception:
        pass
    return "left"


def _safe_line_spacing(pf) -> float:
    try:
        if pf.line_spacing is not None:
            return float(pf.line_spacing)
    except Exception:
        pass
    return 1.0


# ---------------------------------------------------------------------------
# Page geometry
# ---------------------------------------------------------------------------

def _extract_page_geometry(doc) -> tuple[dict[str, float], tuple[float, float]]:
    if not doc.sections:
        return _default_margins(), (8.5, 11.0)

    section = doc.sections[0]
    margins = {
        "top": _to_inches(section.top_margin, 1.0),
        "bottom": _to_inches(section.bottom_margin, 1.0),
        "left": _to_inches(section.left_margin, 1.25),
        "right": _to_inches(section.right_margin, 1.25),
    }
    page_size = (
        _to_inches(section.page_width, 8.5),
        _to_inches(section.page_height, 11.0),
    )
    return margins, page_size


def _default_margins() -> dict[str, float]:
    return {"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25}


def _to_inches(length, default: float) -> float:
    try:
        if length is not None:
            return round(length.inches, 4)
    except Exception:
        pass
    return default


# ---------------------------------------------------------------------------
# Heading hierarchy
# ---------------------------------------------------------------------------

def _extract_heading_hierarchy(doc) -> list[str]:
    """Return unique heading style names in first-appearance order."""
    seen: list[str] = []
    for para in doc.paragraphs:
        name = para.style.name
        if _HEADING_RE.match(name) and name not in seen:
            seen.append(name)
    return seen


# ---------------------------------------------------------------------------
# Placeholder detection
# ---------------------------------------------------------------------------

def _detect_placeholders(doc) -> list[str]:
    """Scan all paragraph text for {{token}} patterns."""
    found: list[str] = []
    for para in doc.paragraphs:
        for match in _PLACEHOLDER_RE.finditer(para.text):
            token = f"{{{{{match.group(1)}}}}}"
            if token not in found:
                found.append(token)
    return found


# ---------------------------------------------------------------------------
# Section skeleton
# ---------------------------------------------------------------------------

def _build_section_skeleton(doc) -> list[dict]:
    """Extract the ordered heading structure of the template."""
    skeleton: list[dict] = []
    for para in doc.paragraphs:
        m = _HEADING_RE.match(para.style.name)
        if m and para.text.strip():
            skeleton.append(
                {
                    "title": para.text.strip(),
                    "style": para.style.name,
                    "level": int(m.group(1)),
                }
            )
    return skeleton
