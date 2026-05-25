"""
Minimal real-pipeline smoke test.

Requires a running Ollama daemon with a compatible model.
Skipped automatically when Ollama is unavailable — safe to run in CI
without a local LLM.

Tests ONLY:
  - plan generation (outline_planner.plan)
  - single-section content generation (content_generator.write_section)
  - DOCX assembly (document_assembler.build)

No PDF export, no UI, no RAG.
"""
import ollama
import pytest
from docx import Document

from app.config import DEFAULT_MODEL, OLLAMA_BASE_URL
from core.content_generator import ContentGeneratorError, write_section
from core.document_assembler import build
from core.models import StyleSpec, TemplateProfile
from core.outline_planner import OutlinePlannerError
from core.outline_planner import plan as plan_report

# ---------------------------------------------------------------------------
# Skip guard
# ---------------------------------------------------------------------------

def _ollama_available() -> bool:
    try:
        ollama.Client(host=OLLAMA_BASE_URL).list()
        return True
    except Exception:
        return False


pytestmark = pytest.mark.skipif(
    not _ollama_available(),
    reason="Ollama daemon not available — skipping integration tests",
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def tiny_template(tmp_path):
    """Minimal DOCX template with all three standard anchors."""
    doc = Document()
    doc.add_paragraph("{{title}}")
    doc.add_paragraph("By {{author}}")
    doc.add_paragraph("{{REPORT_BODY}}")
    path = tmp_path / "tiny.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def tiny_profile():
    """Bare-minimum TemplateProfile: Normal + Heading 1 only."""
    base = StyleSpec(
        name="Normal", font_name="Calibri", font_size_pt=11.0,
        bold=False, italic=False, color_hex=None, alignment="left", line_spacing=1.0,
    )
    heading = StyleSpec(
        name="Heading 1", font_name="Calibri", font_size_pt=14.0,
        bold=True, italic=False, color_hex=None, alignment="left", line_spacing=1.0,
    )
    return TemplateProfile(
        styles={"Normal": base, "Heading 1": heading},
        margins_in={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
        page_size=(8.5, 11.0),
        heading_hierarchy=["Heading 1"],
        placeholders=["{{title}}", "{{author}}", "{{REPORT_BODY}}"],
        section_skeleton=[],
    )


# ---------------------------------------------------------------------------
# Smoke test
# ---------------------------------------------------------------------------

def test_full_pipeline_smoke(tiny_template, tiny_profile, tmp_path, monkeypatch):
    """
    Brief → plan → one section of content → DOCX assembly.
    Validates that the three pipeline stages connect correctly end-to-end.
    """
    monkeypatch.setattr("core.document_assembler.OUTPUTS_DIR", tmp_path)

    # Stage 1: plan
    try:
        report_plan = plan_report(
            brief="A brief introduction to machine learning applications in healthcare.",
            topic="Machine Learning in Healthcare",
            academic_level="undergraduate",
            heading_hierarchy=["Heading 1"],
            model=DEFAULT_MODEL,
        )
    except OutlinePlannerError as exc:
        pytest.skip(f"Plan generation failed (LLM issue): {exc}")

    assert report_plan.title, "Plan must have a non-empty title"
    assert len(report_plan.sections) >= 1, "Plan must have at least one section"

    # Stage 2: content for first section only (keep runtime short)
    first_spec = report_plan.sections[0]
    try:
        section = write_section(
            plan=report_plan,
            section=first_spec,
            topic="Machine Learning in Healthcare",
            previous_summaries=[],
            model=DEFAULT_MODEL,
        )
    except ContentGeneratorError as exc:
        pytest.skip(f"Content generation failed (LLM issue): {exc}")

    assert section.title, "Section must have a non-empty title"
    assert len(section.blocks) >= 1, "Section must contain at least one block"

    # Stage 3: assembly
    out = build(tiny_template, tiny_profile, report_plan, [section])
    assert out.exists(), "Assembled DOCX file must exist"

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(texts) >= 1, "Assembled document must contain text"
    assert "{{REPORT_BODY}}" not in texts, "Anchor must be removed from output"
