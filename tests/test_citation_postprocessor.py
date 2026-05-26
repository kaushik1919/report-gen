"""
Tests for core/citation_postprocessor.py.
All tests use in-memory python-docx Documents — no file I/O required.
"""
from docx import Document

from core.citation_postprocessor import process_document
from core.models import BibliographyStore, CitationEntry

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_bib(*entries: tuple) -> BibliographyStore:
    """Build a BibliographyStore from (key, entry_type, fields) tuples."""
    bib_entries = {}
    for key, entry_type, fields in entries:
        bib_entries[key] = CitationEntry(key=key, entry_type=entry_type, fields=fields)
    return BibliographyStore(entries=bib_entries)


def _sample_bib() -> BibliographyStore:
    return _make_bib(
        ("smith2020", "article", {
            "author": "Smith, John",
            "title": "A Study of ML",
            "journal": "Journal of AI",
            "year": "2020",
        }),
        ("doe2019", "book", {
            "author": "Doe, Jane",
            "title": "Data Science Foundations",
            "year": "2019",
            "publisher": "Academic Press",
        }),
    )


def _doc_with(*paragraphs: str) -> Document:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    return doc


# ---------------------------------------------------------------------------
# TestProcessDocument — marker replacement
# ---------------------------------------------------------------------------

class TestMarkerReplacement:
    def test_replaces_single_marker_with_numbered_ref(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("[1]" in t for t in texts)
        assert not any("REF:" in t for t in texts)

    def test_replaces_marker_without_space_after_colon(self):
        bib = _sample_bib()
        doc = _doc_with("[REF:smith2020]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("[1]" in t for t in texts)

    def test_sequential_numbering_in_appearance_order(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]", "[REF: doe2019]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("[1]" in t for t in texts)
        assert any("[2]" in t for t in texts)

    def test_same_key_gets_same_number_across_paragraphs(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]", "Body text.", "[REF: smith2020]")
        process_document(doc, bib)
        # First 3 paragraphs are the original content (marker, body, marker).
        original_texts = [p.text for p in doc.paragraphs][:3]
        count_1 = sum(t.count("[1]") for t in original_texts)
        assert count_1 == 2

    def test_inline_marker_within_sentence_replaced(self):
        bib = _sample_bib()
        doc = _doc_with("According to [REF: smith2020] this holds true.")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("[1]" in t for t in texts)
        assert any("According to" in t for t in texts)


# ---------------------------------------------------------------------------
# TestProcessDocument — unknown keys
# ---------------------------------------------------------------------------

class TestUnknownKeys:
    def test_unknown_key_preserved_visibly(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: ghost_key]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("ghost_key" in t for t in texts)
        assert any("?" in t for t in texts)

    def test_unknown_key_returns_warning_message(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: ghost_key]")
        warnings = process_document(doc, bib)
        assert len(warnings) == 1
        assert "ghost_key" in warnings[0]

    def test_known_keys_return_no_warnings(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]", "[REF: doe2019]")
        warnings = process_document(doc, bib)
        assert warnings == []

    def test_mixed_known_and_unknown_keys(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]", "[REF: ghost_key]")
        warnings = process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert len(warnings) == 1
        assert any("[1]" in t for t in texts)
        assert any("ghost_key" in t for t in texts)

    def test_multiple_unknown_keys_each_get_warning(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: ghost1]", "[REF: ghost2]")
        warnings = process_document(doc, bib)
        assert len(warnings) == 2


# ---------------------------------------------------------------------------
# TestProcessDocument — bibliography insertion
# ---------------------------------------------------------------------------

class TestBibliographyInsertion:
    def test_bibliography_appended_when_no_anchor(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("References" in t for t in texts)

    def test_bibliography_replaces_anchor_paragraph(self):
        bib = _sample_bib()
        doc = Document()
        doc.add_paragraph("[REF: smith2020]")
        doc.add_paragraph("{{REFERENCES}}")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert "{{REFERENCES}}" not in texts
        assert any("References" in t for t in texts)

    def test_bibliography_entry_appears_after_heading(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: smith2020]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        ref_idx = next(i for i, t in enumerate(texts) if "References" in t)
        assert any("[1]" in t for t in texts[ref_idx + 1:])

    def test_unknown_key_bibliography_entry_has_fallback(self):
        bib = _sample_bib()
        doc = _doc_with("[REF: ghost_key]")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert any("ghost_key" in t for t in texts)


# ---------------------------------------------------------------------------
# TestProcessDocument — edge cases
# ---------------------------------------------------------------------------

class TestEdgeCases:
    def test_empty_document_returns_no_warnings(self):
        bib = _sample_bib()
        doc = Document()
        warnings = process_document(doc, bib)
        assert warnings == []

    def test_no_markers_no_bibliography_section(self):
        bib = _sample_bib()
        doc = _doc_with("No citations in this paragraph.")
        process_document(doc, bib)
        texts = [p.text for p in doc.paragraphs]
        assert not any("References" in t for t in texts)

    def test_empty_bibliography_store_still_processes(self):
        bib = BibliographyStore()
        doc = _doc_with("[REF: smith2020]")
        warnings = process_document(doc, bib)
        assert len(warnings) == 1

    def test_process_does_not_raise_on_complex_document(self):
        bib = _sample_bib()
        doc = Document()
        doc.add_paragraph("Introduction")
        doc.add_paragraph("[REF: smith2020]")
        doc.add_paragraph("Some body text with no citation.")
        doc.add_paragraph("[REF: doe2019]")
        doc.add_paragraph("[REF: smith2020]")
        doc.add_paragraph("{{REFERENCES}}")
        warnings = process_document(doc, bib)
        assert warnings == []
        texts = [p.text for p in doc.paragraphs]
        assert "{{REFERENCES}}" not in texts
