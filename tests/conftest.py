"""
Pytest fixtures that build minimal .docx files programmatically.
No binary fixtures are committed — all templates are generated in tmp_path.
"""
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches


@pytest.fixture
def simple_template(tmp_path: Path) -> Path:
    """
    A DOCX with explicit margins, two heading levels, and {{placeholder}} tokens.
    Used as the primary fixture for style extraction tests.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This report was written by {{author}} on {{date}}.")

    doc.add_heading("Background", level=2)
    doc.add_paragraph("Topic: {{topic}}.")

    doc.add_heading("Methodology", level=1)
    doc.add_paragraph("Standard academic approach.")

    doc.add_heading("Results", level=2)
    doc.add_paragraph("Data collected for {{experiment}}.")

    path = tmp_path / "simple_template.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def minimal_template(tmp_path: Path) -> Path:
    """The most minimal valid DOCX: one paragraph, no headings, no placeholders."""
    doc = Document()
    doc.add_paragraph("Hello world.")
    path = tmp_path / "minimal_template.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def placeholder_only_template(tmp_path: Path) -> Path:
    """DOCX with only placeholder tokens, no headings."""
    doc = Document()
    doc.add_paragraph("Title: {{title}}")
    doc.add_paragraph("Author: {{author}}")
    doc.add_paragraph("Institution: {{institution}}")
    path = tmp_path / "placeholder_only.docx"
    doc.save(str(path))
    return path
