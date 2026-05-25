"""
Tests for document_assembler.build().
All DOCX fixtures are built programmatically — no binaries committed.
"""
from pathlib import Path

import pytest
from docx import Document

from core.document_assembler import build
from core.models import ReportPlan, SectionContent, StyleSpec, TemplateProfile

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _style(name: str) -> StyleSpec:
    return StyleSpec(
        name=name, font_name="Calibri", font_size_pt=11.0,
        bold=False, italic=False, color_hex=None,
        alignment="left", line_spacing=1.0,
    )


def _texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


def _styles(doc: Document) -> list[str]:
    return [p.style.name for p in doc.paragraphs]


def _section(
    title: str = "Introduction",
    level: int = 1,
    blocks: list[dict] | None = None,
) -> SectionContent:
    return SectionContent(
        section_id="sec_01",
        title=title,
        level=level,
        blocks=blocks or [{"type": "paragraph", "text": "Body text."}],
        citations=[],
    )


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def profile() -> TemplateProfile:
    names = ["Heading 1", "Heading 2", "Normal", "List Bullet"]
    return TemplateProfile(
        styles={n: _style(n) for n in names},
        margins_in={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
        page_size=(8.5, 11.0),
        heading_hierarchy=["Heading 1", "Heading 2"],
        placeholders=["{{title}}", "{{author}}"],
        section_skeleton=[],
    )


@pytest.fixture
def plan() -> ReportPlan:
    return ReportPlan(title="Test Report", author="Jane Smith", sections=[])


@pytest.fixture
def template(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("{{title}}")
    doc.add_paragraph("By {{author}}")
    path = tmp_path / "template.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def assemble(monkeypatch, tmp_path):
    """build() wrapper that redirects output to the test's tmp_path."""
    monkeypatch.setattr("core.document_assembler.OUTPUTS_DIR", tmp_path)
    return build


# ---------------------------------------------------------------------------
# Output file
# ---------------------------------------------------------------------------

class TestOutputFile:
    def test_returns_path(self, assemble, template, profile, plan):
        assert isinstance(assemble(template, profile, plan, []), Path)

    def test_output_file_exists(self, assemble, template, profile, plan):
        assert assemble(template, profile, plan, []).exists()

    def test_output_is_valid_docx(self, assemble, template, profile, plan):
        result = assemble(template, profile, plan, [])
        assert Document(str(result)).paragraphs is not None


# ---------------------------------------------------------------------------
# Placeholder replacement
# ---------------------------------------------------------------------------

class TestPlaceholderReplacement:
    def test_title_replaced(self, assemble, template, profile, plan):
        texts = _texts(Document(str(assemble(template, profile, plan, []))))
        assert "Test Report" in texts

    def test_author_replaced(self, assemble, template, profile, plan):
        texts = _texts(Document(str(assemble(template, profile, plan, []))))
        assert any("Jane Smith" in t for t in texts)

    def test_unknown_placeholder_preserved(self, assemble, tmp_path, profile, plan):
        doc = Document()
        doc.add_paragraph("Institution: {{institution}}")
        tpl = tmp_path / "tpl.docx"
        doc.save(str(tpl))
        texts = _texts(Document(str(assemble(tpl, profile, plan, []))))
        assert any("{{institution}}" in t for t in texts)

    def test_cross_run_placeholder_replaced(self, assemble, tmp_path, profile, plan):
        """{{title}} split across two XML runs must still be substituted."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{")
        para.add_run("title}}")
        tpl = tmp_path / "tpl.docx"
        doc.save(str(tpl))
        texts = _texts(Document(str(assemble(tpl, profile, plan, []))))
        assert any("Test Report" in t for t in texts)


# ---------------------------------------------------------------------------
# Heading rendering
# ---------------------------------------------------------------------------

class TestHeadingRendering:
    def test_heading_text_in_output(self, assemble, template, profile, plan):
        texts = _texts(Document(str(assemble(template, profile, plan, [_section("My Heading")]))))
        assert "My Heading" in texts

    def test_heading1_style_applied(self, assemble, template, profile, plan):
        result = assemble(template, profile, plan, [_section(level=1)])
        assert "Heading 1" in _styles(Document(str(result)))

    def test_heading2_style_applied(self, assemble, template, profile, plan):
        result = assemble(template, profile, plan, [_section(level=2)])
        assert "Heading 2" in _styles(Document(str(result)))


# ---------------------------------------------------------------------------
# Paragraph rendering
# ---------------------------------------------------------------------------

class TestParagraphRendering:
    def test_paragraph_text_in_output(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "paragraph", "text": "Academic prose."}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "Academic prose." in texts

    def test_multiple_paragraphs_all_rendered(self, assemble, template, profile, plan):
        section = _section(blocks=[
            {"type": "paragraph", "text": "First."},
            {"type": "paragraph", "text": "Second."},
        ])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "First." in texts
        assert "Second." in texts


# ---------------------------------------------------------------------------
# Bullet rendering
# ---------------------------------------------------------------------------

class TestBulletRendering:
    def test_bullet_items_in_output(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "bullets", "items": ["Alpha", "Beta"]}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "Alpha" in texts
        assert "Beta" in texts

    def test_bullet_count_matches(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "bullets", "items": ["X", "Y", "Z"]}])
        result = assemble(template, profile, plan, [section])
        bullet_paras = [
            p for p in Document(str(result)).paragraphs
            if p.style.name in ("List Bullet", "List Paragraph")
        ]
        assert len(bullet_paras) == 3

    def test_bullet_falls_back_to_list_paragraph(self, assemble, tmp_path, plan):
        """When 'List Bullet' absent from profile, 'List Paragraph' is used instead."""
        prof = TemplateProfile(
            styles={n: _style(n) for n in ["Heading 1", "Normal", "List Paragraph"]},
            margins_in={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
            page_size=(8.5, 11.0),
            heading_hierarchy=["Heading 1"],
            placeholders=[],
            section_skeleton=[],
        )
        tpl = tmp_path / "tpl.docx"
        Document().save(str(tpl))
        section = _section(blocks=[{"type": "bullets", "items": ["Item"]}])
        result = assemble(tpl, prof, plan, [section])
        assert "List Paragraph" in _styles(Document(str(result)))


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

class TestTableRendering:
    def test_table_headers_in_output(self, assemble, template, profile, plan):
        section = _section(blocks=[{
            "type": "table", "headers": ["Name", "Score"], "rows": [],
        }])
        doc = Document(str(assemble(template, profile, plan, [section])))
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "Name"
        assert doc.tables[0].rows[0].cells[1].text == "Score"

    def test_table_data_rows_in_output(self, assemble, template, profile, plan):
        section = _section(blocks=[{
            "type": "table",
            "headers": ["A", "B"],
            "rows": [["1", "2"], ["3", "4"]],
        }])
        doc = Document(str(assemble(template, profile, plan, [section])))
        assert doc.tables[0].rows[1].cells[0].text == "1"
        assert doc.tables[0].rows[2].cells[0].text == "3"

    def test_table_clips_oversized_row(self, assemble, template, profile, plan):
        """Row wider than headers must not raise — extra values are silently dropped."""
        section = _section(blocks=[{
            "type": "table",
            "headers": ["Col"],
            "rows": [["val", "overflow1", "overflow2"]],
        }])
        doc = Document(str(assemble(template, profile, plan, [section])))
        assert doc.tables[0].rows[1].cells[0].text == "val"

    def test_empty_headers_skips_table(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "table", "headers": [], "rows": [["x"]]}])
        doc = Document(str(assemble(template, profile, plan, [section])))
        assert len(doc.tables) == 0


# ---------------------------------------------------------------------------
# Style fallback
# ---------------------------------------------------------------------------

class TestStyleFallback:
    def test_missing_heading_level_falls_back(self, assemble, tmp_path, plan):
        """Heading 3 absent from profile → assembler falls back to Heading 2."""
        prof = TemplateProfile(
            styles={n: _style(n) for n in ["Heading 1", "Heading 2", "Normal"]},
            margins_in={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
            page_size=(8.5, 11.0),
            heading_hierarchy=["Heading 1", "Heading 2"],
            placeholders=[],
            section_skeleton=[],
        )
        tpl = tmp_path / "tpl.docx"
        Document().save(str(tpl))
        result = assemble(tpl, prof, plan, [_section(level=3)])
        assert "Heading 2" in _styles(Document(str(result)))

    def test_unknown_block_type_warns_and_skips(self, assemble, template, profile, plan):
        """Unrecognised block type logs a WARNING and is skipped; subsequent blocks render."""
        section = _section(blocks=[
            {"type": "figure", "caption": "Fig 1.", "image_ref": "fig1.png"},
            {"type": "paragraph", "text": "After figure."},
        ])
        assert "After figure." in _texts(Document(str(assemble(template, profile, plan, [section]))))


# ---------------------------------------------------------------------------
# Anchor insertion
# ---------------------------------------------------------------------------

@pytest.fixture
def template_with_anchor(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Preamble text.")
    doc.add_paragraph("{{REPORT_BODY}}")
    doc.add_paragraph("Postamble text.")
    path = tmp_path / "anchored.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def template_with_all_anchors(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("{{REPORT_BODY}}")
    doc.add_paragraph("{{REFERENCES}}")
    doc.add_paragraph("{{APPENDIX}}")
    path = tmp_path / "all_anchors.docx"
    doc.save(str(path))
    return path


class TestAnchorInsertion:
    def test_report_body_anchor_removed(self, assemble, template_with_anchor, profile, plan):
        """{{REPORT_BODY}} paragraph must not appear in the output."""
        texts = _texts(Document(str(assemble(template_with_anchor, profile, plan, []))))
        assert "{{REPORT_BODY}}" not in texts

    def test_sections_inserted_at_anchor_position(
        self, assemble, template_with_anchor, profile, plan
    ):
        """Section heading appears between preamble and postamble."""
        texts = _texts(Document(str(
            assemble(template_with_anchor, profile, plan, [_section("My Heading")])
        )))
        preamble_idx = texts.index("Preamble text.")
        heading_idx = texts.index("My Heading")
        postamble_idx = texts.index("Postamble text.")
        assert preamble_idx < heading_idx < postamble_idx

    def test_fallback_appends_when_no_anchor(self, assemble, template, profile, plan):
        """Template without {{REPORT_BODY}}: sections appended after template content."""
        texts = _texts(Document(str(
            assemble(template, profile, plan, [_section("Appended Section")])
        )))
        assert "Appended Section" in texts

    def test_references_anchor_cleaned_up(
        self, assemble, template_with_all_anchors, profile, plan
    ):
        """{{REFERENCES}} anchor is always removed from output (reserved for M4)."""
        texts = _texts(Document(str(assemble(template_with_all_anchors, profile, plan, []))))
        assert "{{REFERENCES}}" not in texts

    def test_appendix_anchor_cleaned_up(
        self, assemble, template_with_all_anchors, profile, plan
    ):
        """{{APPENDIX}} anchor is always removed from output (reserved for M4)."""
        texts = _texts(Document(str(assemble(template_with_all_anchors, profile, plan, []))))
        assert "{{APPENDIX}}" not in texts

    def test_anchor_table_inserted_at_position(
        self, assemble, template_with_anchor, profile, plan
    ):
        """Table block within an anchored section ends up before postamble."""
        section = _section(blocks=[{
            "type": "table", "headers": ["Col"], "rows": [["val"]],
        }])
        doc = Document(str(assemble(template_with_anchor, profile, plan, [section])))
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "Col"


# ---------------------------------------------------------------------------
# New block types
# ---------------------------------------------------------------------------

class TestNewBlockTypes:
    def test_bullet_list_type_renders(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "bullet_list", "items": ["Alpha", "Beta"]}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "Alpha" in texts
        assert "Beta" in texts

    def test_numbered_list_renders(self, assemble, template, profile, plan):
        prof = TemplateProfile(
            styles={n: _style(n) for n in ["Heading 1", "Normal", "List Number"]},
            margins_in={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
            page_size=(8.5, 11.0),
            heading_hierarchy=["Heading 1"],
            placeholders=[],
            section_skeleton=[],
        )
        section = _section(blocks=[{"type": "numbered_list", "items": ["One", "Two"]}])
        texts = _texts(Document(str(assemble(template, prof, plan, [section]))))
        assert "One" in texts
        assert "Two" in texts

    def test_inline_heading_renders(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "heading", "text": "Sub-heading", "level": 2}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "Sub-heading" in texts

    def test_figure_placeholder_renders(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "figure_placeholder", "caption": "Model diagram"}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert any("Figure" in t and "Model diagram" in t for t in texts)

    def test_citation_placeholder_renders(self, assemble, template, profile, plan):
        section = _section(blocks=[{"type": "citation_placeholder", "key": "smith2020"}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert any("REF" in t and "smith2020" in t for t in texts)

    def test_bullets_legacy_alias_still_works(self, assemble, template, profile, plan):
        """Legacy 'bullets' type must remain functional."""
        section = _section(blocks=[{"type": "bullets", "items": ["Item"]}])
        texts = _texts(Document(str(assemble(template, profile, plan, [section]))))
        assert "Item" in texts
